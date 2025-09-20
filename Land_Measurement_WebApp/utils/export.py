from openpyxl import load_workbook
from openpyxl.drawing.image import Image
from io import BytesIO
from .visualization import make_graphic_file_io
from docx import Document

def export_to_excel(table):
    '''Export survey data to Excel with embedded graphs.
    
    Creates an Excel file containing:
    - Survey data table
    - Elevation profile graph
    - Height difference graph
    - Elevation scatter plot
    
    Args:
        table: Pandas DataFrame containing survey data
        file_path: Path to save the Excel file
    '''
   
    buffers=make_graphic_file_io(table)
    excel_buffer = BytesIO()
    table.to_excel(excel_buffer,float_format='%.3f',index=False)
    excel_buffer.seek(0)
    wb = load_workbook(excel_buffer)
    ws = wb.active
    def add_image_to_excel(buffer, cell):
        no_image = Image(buffer)
        ws.add_image(no_image,cell)
    add_image_to_excel(buffers['elevation_profile'], 'I2')
    add_image_to_excel(buffers['height_difference'], 'I30')
    add_image_to_excel(buffers['elevation_scatter'], 'I58')
    return excel_buffer
    

        

def export_to_csv(table,file_path):
    '''Export survey data to CSV file.
    
    Args:
        table: Pandas DataFrame containing survey data
        file_path: Path to save the CSV file
    '''
    
    table.to_csv(file_path, index=False, float_format="%.3f")
    
    


def export_to_word(report_text, filename="document.docx"):
    doc = Document()
    doc.add_heading("Survey Analysis Report", 0)
    # Pisahkan baris biar rapi
    for line in report_text.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")  # baris kosong
        elif line.startswith("**"):  # heading bold
            doc.add_heading(line.replace("**", ""), level=1)
        else:
            doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer